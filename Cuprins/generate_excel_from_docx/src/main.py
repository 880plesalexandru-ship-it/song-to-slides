import os
import re
from docx import Document
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

from docx.oxml.shared import OxmlElement, qn

def genereaza_docx(cuprins, fisier_output):
    """
    Generează un fișier Word structurat corect cu cântările din cuprins.
    Structura:
    - Cuprins alfabetic (pe două coloane)
    - Cuprins tematic (pe două coloane)
    """
    document = Document()

    # Adăugăm titlul principal
    document.add_heading("Cuprins", level=1)

    # Secțiunea "Cuprins Alfabetic"
    document.add_heading("Cuprins Alfabetic", level=2)
    # Filtrăm cântările care nu sunt de la nuntă sau colinde și sortăm alfabetic după nume
    cantarile_alfabetice = sorted(
        [cantare for cantare in cuprins if cantare["Tematica"] not in ["NUNTĂ", "COLINDE"]],
        key=lambda x: x["Nume"].lower()
    )

    # Organizăm cântările în două coloane
    tabel_alfabetic = document.add_table(rows=1, cols=2)
    tabel_alfabetic.style = 'Table Grid'

    # Setăm bordurile tabelului la alb
    tbl = tabel_alfabetic._tbl
    tbl_pr = tbl.tblPr  # Proprietățile tabelului
    tbl_borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # Stilul bordurii
        border.set(qn('w:sz'), '4')  # Grosimea bordurii
        border.set(qn('w:color'), 'FFFFFF')  # Culoarea albă
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)

    litera_curenta = None
    coloana_stanga = []
    coloana_dreapta = []
    ultima_coloana = None  # Păstrăm ultima coloană utilizată (stânga sau dreapta)

    for index, cantare in enumerate(cantarile_alfabetice):
        # Eliminăm simbolurile "/:" și ’ de la începutul numelui cântării
        cantare["Nume"] = cantare["Nume"].lstrip("/:’")

        prima_litera = cantare["Nume"][0].upper()
        if prima_litera != litera_curenta:

            litera_curenta = prima_litera
            # Adăugăm o linie goală sub litera curentă
            coloana_stanga.append("")
            coloana_dreapta.append("")
            # Adăugăm litera în fiecare coloană
            coloana_stanga.append(f"{prima_litera}")
            coloana_dreapta.append("")

        # Adăugăm cântarea în coloana corespunzătoare
        if index % 2 == 0:
            coloana_stanga.append(f"{cantare['Numar']}\t{cantare['Nume']}")
            ultima_coloana = "stanga"
        else:
            coloana_dreapta.append(f"{cantare['Numar']}\t{cantare['Nume']}")
            ultima_coloana = "dreapta"

    # Verificăm dacă ultima cântare a fost în coloana stângă
    if ultima_coloana == "stanga":
        coloana_dreapta.append("")  # Adăugăm o linie goală în coloana dreaptă

    # Adăugăm cântările în tabel
    max_len = max(len(coloana_stanga), len(coloana_dreapta))
    for i in range(max_len):
        stanga = coloana_stanga[i] if i < len(coloana_stanga) else ""
        dreapta = coloana_dreapta[i] if i < len(coloana_dreapta) else ""
        row = tabel_alfabetic.add_row().cells
    
        # Adăugăm textul în coloana stângă
        if stanga:
            if len(stanga) == 1:  # Verificăm dacă textul are o singură literă
                run_stanga = row[0].paragraphs[0].add_run(stanga)
                run_stanga.bold = True  # Aplicăm stilul bold
            else:
                row[0].text = stanga
    
        # Adăugăm textul în coloana dreaptă
        if dreapta:
            if len(dreapta) == 1:  # Verificăm dacă textul are o singură literă
                run_dreapta = row[1].paragraphs[0].add_run(dreapta)
                run_dreapta.bold = True  # Aplicăm stilul bold
            else:
                row[1].text = dreapta

    # Secțiunea "Cuprins Tematic"
    document.add_heading("Cuprins Tematic", level=2)
    tematici = {}
    for cantare in cuprins:
        if cantare["Tematica"]:  # Excludem cântările fără tematică
            if cantare["Tematica"] not in tematici:
                tematici[cantare["Tematica"]] = []
            tematici[cantare["Tematica"]].append(cantare)

    # Sortăm tematicile alfabetic
    tematici_sortate = sorted(tematici.items(), key=lambda x: x[0].lower())

    for tematica, cantarile in tematici_sortate:
        document.add_heading(tematica, level=3)
        tabel_tematic = document.add_table(rows=1, cols=2)
        tabel_tematic.style = 'Table Grid'

        # Setăm bordurile tabelului la alb
        tbl = tabel_tematic._tbl
        tbl_pr = tbl.tblPr
        tbl_borders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), 'FFFFFF')
            tbl_borders.append(border)

        tbl_pr.append(tbl_borders)

        coloana_stanga = []
        coloana_dreapta = []
        for index, cantare in enumerate(cantarile):
            if index % 2 == 0:
                coloana_stanga.append(f"{cantare['Numar']}\t{cantare['Nume']}")
            else:
                coloana_dreapta.append(f"{cantare['Numar']}\t{cantare['Nume']}")

        max_len = max(len(coloana_stanga), len(coloana_dreapta))
        for i in range(max_len):
            stanga = coloana_stanga[i] if i < len(coloana_stanga) else ""
            dreapta = coloana_dreapta[i] if i < len(coloana_dreapta) else ""
            row = tabel_tematic.add_row().cells
            row[0].text = stanga
            row[1].text = dreapta

    # Salvăm documentul
    document.save(fisier_output)
    print(f"Fișierul Word '{fisier_output}' a fost creat cu succes!")


def citeste_docx(fisier):
    document = Document(fisier)
    l_cantarile = []
    tematica_curenta = None
    sectiune_curenta = None

    for par in document.paragraphs:
        text = par.text.strip()
        if not text:
            continue

        # Detectăm secțiunea curentă
        if text.lower() == "alfabetic":
            sectiune_curenta = "alfabetic"
            continue
        elif text.lower() == "tematic":
            sectiune_curenta = "tematic"
            continue
        elif text.startswith("CUPRINS – NUNTĂ"):
            sectiune_curenta = "nunta"
            tematica_curenta = "NUNTĂ"
            continue
        elif text.startswith("CUPRINS – COLINDE"):
            sectiune_curenta = "colinde"
            tematica_curenta = "COLINDE"
            continue

        # Dacă suntem în secțiunea tematică, detectăm tematica
        if sectiune_curenta == "tematic" and text.isupper():
            tematica_curenta = text
            continue

        # Potrivim formatul "Titlu – Număr"
        match = re.match(r"^(.+?)\s+–\s+(\d+)$", text)
        if match:
            titlu = match.group(1).strip()
            numar = int(match.group(2).strip())

            # Curățăm titlul
            titlu = titlu.lstrip("/:’ ").rstrip(",.")

            # Dacă suntem în secțiunea "alfabetic", adăugăm cântarea
            if sectiune_curenta == "alfabetic":
                l_cantarile.append({
                    "Numar": numar,
                    "Nume": titlu,
                    "Tematica": "",  # Tematica va fi completată ulterior
                    "Gama": ""  # Gama va fi completată ulterior
                })

            # Dacă suntem în secțiunea "tematic", actualizăm tematica cântării
            elif sectiune_curenta == "tematic" and tematica_curenta:
                for cantare in l_cantarile:
                    if cantare["Nume"] == titlu:
                        cantare["Tematica"] = tematica_curenta
                        break

            # Dacă suntem în secțiunile speciale "nunta" sau "colinde", adăugăm cântarea
            elif sectiune_curenta in ["nunta", "colinde"]:
                l_cantarile.append({
                    "Numar": numar,
                    "Nume": titlu,
                    "Tematica": tematica_curenta,  # Setăm tematica specifică secțiunii
                    "Gama": ""  # Gama va fi completată ulterior
                })

    return l_cantarile

def citeste_caiet(fisier):
    """
    Citește cântările din fișierul caiet.docx în formatul:
    - Numărul urmat de un punct pe un rând (ex. "32.")
    - Titlul pe rândul următor
    """
    document = Document(fisier)
    l_cantarile = []
    numar_curent = None

    for par in document.paragraphs:
        text = par.text.strip()
        if not text:
            continue

        # Verificăm dacă textul este un număr urmat de un punct (ex. "32.")
        match = re.match(r"^(\d+)\.$", text)
        if match:
            numar_curent = int(match.group(1))  # Salvăm numărul curent
        elif numar_curent is not None:
            # Considerăm că acest rând este titlul cântării
            titlu = text.lstrip("/:’ ").rstrip(",.")  # Eliminăm simbolurile doar dacă sunt la început
            # Limităm titlul la maximum 4 cuvinte
            titlu = " ".join(titlu.split()[:4])
            l_cantarile.append({
                "number": numar_curent,
                "name": titlu
            })
            numar_curent = None  # Resetăm numărul curent pentru următoarea cântare
    
    return l_cantarile

def genereaza_excel(cantarile, fisier_output):
    df = pd.DataFrame(cantarile)
    df.to_excel(fisier_output, index=False, sheet_name='Cantarile')
    
    # Adăugăm filtre pe coloane
    with pd.ExcelWriter(fisier_output, engine='openpyxl', mode='a') as writer:
        workbook = writer.book
        worksheet = writer.sheets['Cantarile']
        worksheet.auto_filter.ref = worksheet.dimensions

def completeaza_cantarile(cuprins, caiet):
    """
    Completează cântările lipsă din cuprins folosind cântările din caiet.
    """
    # Creăm un set cu numerele cântărilor din cuprins pentru verificare rapidă
    numere_cuprins = {cantare["Numar"] for cantare in cuprins}

    # Adăugăm cântările din caiet care lipsesc în cuprins
    for cantare in caiet:
        if cantare["number"] not in numere_cuprins:
            cuprins.append({
                "Numar": cantare["number"],
                "Nume": cantare["name"],
                "Tematica": "",  # Nu avem tematică pentru cântările lipsă
                "Gama": ""  # Gama va fi completată ulterior
            })

    # Sortăm lista după număr pentru a păstra ordinea
    cuprins.sort(key=lambda x: x["Numar"])

    return cuprins

def main():
    fisier_cuprins = "C:/Users/aplesa/Desktop/Personal Files/Adunare/Tineret/PPT_Slides/python/Cuprins/generate_excel_from_docx/Cuprins_nou.docx"
    fisier_caiet = "C:/Users/aplesa/Desktop/Personal Files/Adunare/Tineret/PPT_Slides/python/caiet.docx"
    fisier_excel = "C:/Users/aplesa/Desktop/Personal Files/Adunare/Tineret/PPT_Slides/python/Cuprins/generate_excel_from_docx/cantarile.xlsx"
    fisier_docx = "C:/Users/aplesa/Desktop/Personal Files/Adunare/Tineret/PPT_Slides/python/Cuprins/generate_excel_from_docx/Cuprins_structurat.docx"

    # Citim cântările din cuprins și caiet
    cuprins = citeste_docx(fisier_cuprins)
    caiet = citeste_caiet(fisier_caiet)

    # Completăm cântările lipsă
    cuprins_complet = completeaza_cantarile(cuprins, caiet)

    # Generăm fișierul Excel
    genereaza_excel(cuprins_complet, fisier_excel)

    # Generăm fișierul Word structurat
    genereaza_docx(cuprins_complet, fisier_docx)


if __name__ == "__main__":
    main()