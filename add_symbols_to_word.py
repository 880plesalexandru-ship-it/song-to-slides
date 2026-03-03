"""
Script pentru a adăuga simboluri de ordine (S1, S2, R1, R2, etc.) în documentul Word
bazat pe detectarea automată a strofelor și refrenelor.

Editează documentul in-place, păstrând fontul, dimensiunea și alinierea originală.
Creează automat un backup înainte de modificare.

Utilizare:
    python add_symbols_to_word.py

Reguli de ordine:
  - Un singur refren → notat R (fără număr), repetat după fiecare strofă:
      S1 R S2 R S3 R ...
  - Mai multe refrene → R1 se repetă (nr_strofe - nr_refrene) ori extra,
      apoi R2, R3... apar câte o dată:
      S1 R1 S2 R1 S3 R2

Formatul rezultat:
    24. | S1 R S2 R S3 R
    S1  Eu pe Domnul iubesc...        (S1 bold, aliniat la stânga)
    Pentru că m-a creat...
        R  Cerul, pământul...          (R bold, înainte de tab)
        Să-I cânte toată suflarea...
"""

import re
import shutil
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _este_refren(paragraf):
    """Verifică dacă un paragraf este indentat (tab sau riglă)."""
    text = "".join(run.text for run in paragraf.runs)
    indent_prima = paragraf.paragraph_format.first_line_indent
    indent_stanga = paragraf.paragraph_format.left_indent

    return (
        text.startswith("\t") or
        (indent_prima and indent_prima > Inches(0)) or
        (indent_stanga and indent_stanga > Inches(0))
    )


def _proceseaza_bloc(linii, elemente):
    """
    Procesează un bloc de linii consecutive (fără linii goale).
    Dacă blocul conține o tranziție non-tab → tab, se împarte în strofă + refren.
    Dacă tabul apare de la început, totul e refren.
    Dacă nu apare tab deloc, totul e strofă.

    Liniile tab→non-tab NU creează elemente noi (sunt linii de refren
    cărora le lipsește tab-ul din formatarea originală).
    """
    if not linii:
        return

    # Găsește prima linie cu tab/indent
    primul_refren = None
    for i, (idx, is_ref) in enumerate(linii):
        if is_ref:
            primul_refren = i
            break

    if primul_refren is None:
        # Tot blocul e strofă
        elemente.append({"tip": "strofa", "prima_linie": linii[0][0]})
    elif primul_refren == 0:
        # Tot blocul e refren (începe cu tab)
        elemente.append({"tip": "refren", "prima_linie": linii[0][0]})
    else:
        # Strofă apoi refren
        elemente.append({"tip": "strofa", "prima_linie": linii[0][0]})
        elemente.append({"tip": "refren", "prima_linie": linii[primul_refren][0]})


def _genereaza_ordine_display(elemente):
    """
    Generează ordinea de afișare (performanță) cu reguli de intercalare:
      - 0 refrene → doar strofe
      - 1 refren → R (fără număr), repetat după fiecare strofă
      - mai multe refrene, dar mai puține decât strofele →
          R1 apare (nr_strofe - nr_refrene + 1) ori, apoi R2, R3... câte o dată
      - mai multe/egale refrene vs strofe → ordinea fizică
    """
    stanzas = [e for e in elemente if e["tip"] == "strofa"]
    choruses = [e for e in elemente if e["tip"] == "refren"]
    num_s = len(stanzas)
    num_r = len(choruses)

    if num_r == 0:
        return [f"S{i+1}" for i in range(num_s)]

    if num_s == 0:
        if num_r == 1:
            return ["R"]
        return [f"R{i+1}" for i in range(num_r)]

    starts_with_chorus = elemente[0]["tip"] == "refren"

    if num_r == 1:
        # Singur refren → "R" după fiecare strofă
        ordine = []
        if starts_with_chorus:
            ordine.append("R")
        for i in range(num_s):
            ordine.append(f"S{i+1}")
            ordine.append("R")
        return ordine

    if num_r >= num_s:
        # Mai multe refrene decât strofe → ordinea fizică
        idx_s = idx_r = 0
        ordine = []
        for e in elemente:
            if e["tip"] == "strofa":
                idx_s += 1
                ordine.append(f"S{idx_s}")
            else:
                idx_r += 1
                ordine.append(f"R{idx_r}")
        return ordine

    # Mai multe strofe decât refrene → R1 se repetă
    r1_total = num_s - num_r + 1
    ordine = []
    if starts_with_chorus:
        ordine.append("R1")
        r1_total = max(0, r1_total - 1)

    r1_used = 0
    r_next = 2
    for i in range(num_s):
        ordine.append(f"S{i+1}")
        if r1_used < r1_total:
            ordine.append("R1")
            r1_used += 1
        elif r_next <= num_r:
            ordine.append(f"R{r_next}")
            r_next += 1

    return ordine


def _genereaza_prefixuri(elemente):
    """
    Generează etichetele prefix pentru fiecare element.
    Un singur refren → "R"; mai multe → "R1", "R2" etc.
    """
    choruses = [e for e in elemente if e["tip"] == "refren"]
    num_r = len(choruses)

    prefixuri = {}
    idx_s = idx_r = 0
    for elem in elemente:
        if elem["tip"] == "strofa":
            idx_s += 1
            prefixuri[elem["prima_linie"]] = {
                "prefix": f"S{idx_s}", "is_refren": False
            }
        else:
            idx_r += 1
            label = "R" if num_r == 1 else f"R{idx_r}"
            prefixuri[elem["prima_linie"]] = {
                "prefix": label, "is_refren": True
            }
    return prefixuri


def _inserare_prefix_bold(paragraf, prefix, is_refren):
    """
    Inserează prefix bold (S1, R etc.) ca run separat înaintea textului existent.
    - Pentru strofe: adaugă negative first_line_indent ca prefixul să iasă în stânga
    - Pentru refrene cu tab: pune prefixul înainte de tab, textul rămâne aliniat
    """
    if not paragraf.runs:
        return

    first_run_elem = paragraf.runs[0]._r
    p_elem = paragraf._p

    # Creează run bold nou
    new_r = OxmlElement('w:r')
    new_rPr = OxmlElement('w:rPr')
    new_rPr.append(OxmlElement('w:b'))

    # Copiază proprietățile fontului din primul run existent
    existing_rPr = first_run_elem.find(qn('w:rPr'))
    if existing_rPr is not None:
        for tag in ['w:rFonts', 'w:sz', 'w:szCs', 'w:color']:
            prop = existing_rPr.find(qn(tag))
            if prop is not None:
                new_rPr.append(deepcopy(prop))
    new_r.append(new_rPr)

    # Creează elementul text
    new_t = OxmlElement('w:t')
    new_t.set(qn('xml:space'), 'preserve')

    has_tab = paragraf.runs[0].text.startswith('\t')

    if is_refren and has_tab:
        # Pune prefixul înainte de tab: "R\tText" → prefixul la stânga, textul la tab-stop
        paragraf.runs[0].text = paragraf.runs[0].text[1:]  # Elimină tab-ul
        new_t.text = prefix + '\t'
    else:
        new_t.text = prefix + ' '

    new_r.append(new_t)

    # Inserează noul run înaintea primului run existent
    p_elem.insert(list(p_elem).index(first_run_elem), new_r)

    # Indent negativ pentru paragrafele fără tab (strofă sau refren cu indent)
    if not (is_refren and has_tab):
        indent_pts = len(prefix) * 7 + 5
        paragraf.paragraph_format.first_line_indent = -Pt(indent_pts)


def _analizeaza_structura(document):
    """
    Analizează documentul și returnează un dict {index_paragraf: acțiune}.

    Acțiunile pot fi:
      - {"tip": "numar", "ordine": ["S1", "R", ...]}
      - {"tip": "prefix", "prefix": "S1", "is_refren": False}
    """
    paragrafe = document.paragraphs
    actiuni = {}

    # Stare
    in_cantare = False
    numar_idx = None
    elemente = []
    bloc_curent = []  # [(index_paragraf, is_refren)]

    def _finalizeaza_cantare():
        """Salvează acțiunile pentru cântarea curentă."""
        nonlocal elemente, bloc_curent
        _proceseaza_bloc(bloc_curent, elemente)
        bloc_curent = []

        if not elemente:
            return

        ordine = _genereaza_ordine_display(elemente)
        actiuni[numar_idx] = {"tip": "numar", "ordine": ordine}

        prefixuri = _genereaza_prefixuri(elemente)
        for idx_para, info in prefixuri.items():
            actiuni[idx_para] = {
                "tip": "prefix",
                "prefix": info["prefix"],
                "is_refren": info["is_refren"],
            }

    for idx, paragraf in enumerate(paragrafe):
        text = "".join(run.text for run in paragraf.runs).rstrip()

        # Linie goală → finalizează blocul curent
        if not text:
            if bloc_curent:
                _proceseaza_bloc(bloc_curent, elemente)
                bloc_curent = []
            continue

        # Număr de cântare (ex: "24.")
        match = re.match(r"^(\d+)\.$", text)
        if match:
            # Finalizează cântarea anterioară
            if in_cantare:
                _finalizeaza_cantare()

            numar_idx = idx
            elemente = []
            bloc_curent = []
            in_cantare = True
            continue

        if not in_cantare:
            continue

        # Adaugă linia la blocul curent
        bloc_curent.append((idx, _este_refren(paragraf)))

    # Finalizează ultima cântare
    if in_cantare:
        _finalizeaza_cantare()

    return actiuni


def proceseaza_document(fisier, tonalitati=None):
    """
    Editează documentul in-place, adăugând prefixe S1/R1 și ordinea pe linia numărului.

    Args:
        fisier: Calea documentului .docx
        tonalitati: Dict opțional {numar_cantare: "tonalitate"} pentru a adăuga
                    tonalitatea pe linia cu numărul (ex: {24: "Re"})
    """
    # Backup
    backup = fisier.replace(".docx", "_backup.docx")
    shutil.copy2(fisier, backup)
    print(f"Backup creat: {backup}")

    document = Document(fisier)
    actiuni = _analizeaza_structura(document)

    modificari = 0
    for idx, paragraf in enumerate(document.paragraphs):
        if idx not in actiuni:
            continue

        actiune = actiuni[idx]

        if actiune["tip"] == "numar":
            text = "".join(run.text for run in paragraf.runs).rstrip()
            ordine_str = " ".join(actiune["ordine"])

            match = re.match(r"^(\d+)\.$", text)
            if match and paragraf.runs:
                numar = int(match.group(1))
                if tonalitati and numar in tonalitati:
                    paragraf.runs[0].text = f"{text} | {tonalitati[numar]} | {ordine_str}"
                else:
                    paragraf.runs[0].text = f"{text} | {ordine_str}"
                modificari += 1

        elif actiune["tip"] == "prefix":
            _inserare_prefix_bold(
                paragraf, actiune["prefix"], actiune["is_refren"]
            )
            modificari += 1

    document.save(fisier)
    print(f"Document salvat: {fisier}")
    print(f"Total modificări: {modificari}")
    print("Documentul a fost modificat cu succes!")


if __name__ == "__main__":
    proceseaza_document("caiet.docx")
